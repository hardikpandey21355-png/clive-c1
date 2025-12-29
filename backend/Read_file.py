import PyPDF2
import docx
from PIL import Image
import pytesseract
import os
import fitz  # PyMuPDF
import io

# ✅ Windows Tesseract path
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

def read_pdf(file_path):
    """
    ✅ ENHANCED: Reads both TEXT and IMAGES from PDF
    """
    try:
        text = ""
        
        # Step 1: Extract regular text using PyPDF2
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        
        # Step 2: Extract images and OCR them using PyMuPDF
        pdf_document = fitz.open(file_path)
        
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            
            # Get all images on this page
            image_list = page.get_images(full=True)
            
            if image_list:
                text += f"\n--- Images found on page {page_num + 1} ---\n"
            
            for img_index, img in enumerate(image_list):
                try:
                    # Extract image
                    xref = img[0]
                    base_image = pdf_document.extract_image(xref)
                    image_bytes = base_image["image"]
                    
                    # Convert to PIL Image
                    image = Image.open(io.BytesIO(image_bytes))
                    
                    # OCR the image
                    image_text = pytesseract.image_to_string(image)
                    
                    if image_text.strip():
                        text += f"\n[Image {img_index + 1} content]:\n{image_text}\n"
                
                except Exception as img_error:
                    print(f"⚠️ Could not extract image {img_index + 1}: {str(img_error)}")
                    continue
        
        pdf_document.close()
        
        return text.strip() if text.strip() else "No text or readable content found in PDF"
    
    except Exception as e:
        return f"Error reading PDF: {str(e)}"


def read_docx(file_path):
    """
    ✅ ENHANCED: Reads text AND images from DOCX
    """
    try:
        doc = docx.Document(file_path)
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                text += row_text + "\n"
        
        # Extract images (if any embedded)
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    image_data = rel.target_part.blob
                    image = Image.open(io.BytesIO(image_data))
                    image_text = pytesseract.image_to_string(image)
                    if image_text.strip():
                        text += f"\n[Image content]:\n{image_text}\n"
                except:
                    continue
        
        return text.strip() if text.strip() else "No content found in DOCX"
    
    except Exception as e:
        return f"Error reading DOCX: {str(e)}"


def read_txt(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read().strip()
    except Exception as e:
        return f"Error reading TXT: {str(e)}"


def read_image(file_path):
    """
    ✅ Reads text from images (screenshots, photos, scanned docs)
    """
    try:
        image = Image.open(file_path)
        
        # Enhance image for better OCR
        image = image.convert('L')  # Convert to grayscale
        
        # OCR with better config
        custom_config = r'--oem 3 --psm 6'
        text = pytesseract.image_to_string(image, config=custom_config)
        
        return text.strip() if text.strip() else "No text found in image"
    
    except Exception as e:
        return f"Error reading image: {str(e)}"


def read_file(file_path):
    """
    ✅ Main function - routes to correct reader
    """
    if not os.path.exists(file_path):
        return "File not found"
    
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    
    print(f"📄 Reading file: {file_path} (type: {ext})")
    
    if ext == '.pdf':
        return read_pdf(file_path)
    elif ext == '.docx':
        return read_docx(file_path)
    elif ext == '.txt':
        return read_txt(file_path)
    elif ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp']:
        return read_image(file_path)
    else:
        return f"Unsupported file type: {ext}"